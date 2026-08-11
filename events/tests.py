from datetime import timedelta
from io import BytesIO
from zipfile import ZipFile

from django.core.exceptions import ValidationError
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase
from django.urls import reverse
from django.utils import timezone
from docx import Document
from openpyxl import Workbook
from PIL import Image

from accounts.models import User
from events.models import (
    Event,
    EventCategory,
    SpecialEventParticipant,
    SpecialEventPublication,
)
from events.services import import_special_event_participants


class PublicNavigationTests(TestCase):
    def test_public_navigation_links_to_registration_status(self):
        response = self.client.get("/en/")
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Registration Status")
        self.assertContains(response, "/en/registration-status/")
        self.assertContains(response, "← Back")
        self.assertContains(response, "Exit")

    def test_staff_modules_are_displayed_below_available_events(self):
        user = User.objects.create_user(
            username="event.navigation",
            password="Strong-Test-Password-2026",
            role=User.Role.EVENT_ADMIN,
            preferred_language=User.PreferredLanguage.ENGLISH,
        )
        self.client.force_login(user)

        response = self.client.get("/en/")

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="available-events"')
        self.assertContains(response, 'class="module-launcher"')
        self.assertContains(response, "Exhibitions")
        self.assertContains(response, "/en/staff/meetings/")
        self.assertNotContains(response, "meetings-link")


class SpecialEventParticipantQRTests(TestCase):
    password = "Strong-Test-Password-2026"

    def setUp(self):
        self.user = User.objects.create_user(
            username="special.event.admin",
            email="special-admin@example.test",
            password=self.password,
            role=User.Role.EVENT_ADMIN,
            preferred_language=User.PreferredLanguage.ENGLISH,
        )
        self.special_category = EventCategory.objects.create(
            code="SPECIAL_EVENT",
            name_sw="Tukio Maalum",
            name_en="Special Event",
            slug="special-event",
        )
        self.other_category = EventCategory.objects.create(
            code="TRAINING",
            name_sw="Mafunzo",
            name_en="Training",
            slug="training",
        )
        now = timezone.now()
        self.event = Event.objects.create(
            category=self.special_category,
            code="WATAFITI-2026",
            title_sw="Tukio la Watafiti",
            title_en="Researchers Special Event",
            starts_at=now + timedelta(days=2),
            ends_at=now + timedelta(days=3),
            status=Event.Status.PUBLISHED,
            is_public=True,
        )
        self.other_event = Event.objects.create(
            category=self.other_category,
            code="TRAINING-2026",
            title_sw="Mafunzo",
            title_en="Training",
            starts_at=now + timedelta(days=2),
            ends_at=now + timedelta(days=3),
        )

    def workbook_upload(self, rows_by_sheet=None):
        rows_by_sheet = rows_by_sheet or {
            "AWAMU_2022_23": [
                (1, "Researcher One", "Institution A", "Research A", "Agriculture", "2022_2023"),
                (2, "Repeated Name", "Institution B", "Research B", "Health", "2022_2023"),
            ],
            "AWAMU_2023_24": [
                (1, "Repeated Name", "Institution C", "Research C", "Education", "2023_2024"),
            ],
        }
        workbook = Workbook()
        workbook.remove(workbook.active)
        for title, rows in rows_by_sheet.items():
            sheet = workbook.create_sheet(title)
            number_header = "Na." if title == "AWAMU_2023_24" else "Na"
            sheet.append([
                number_header,
                "NAME",
                "INSTITUTION",
                "RESEARCH TITLE",
                "AWARD CATEGORY",
                "YEAR",
            ])
            for row in rows:
                sheet.append(row)
        output = BytesIO()
        workbook.save(output)
        workbook.close()
        return SimpleUploadedFile(
            "researchers.xlsx",
            output.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    def create_researcher(self, *, name="Researcher", institution="Institution"):
        return SpecialEventParticipant.objects.create(
            event=self.event,
            full_name=name,
            institution=institution,
        )

    def create_publication(
        self,
        participant,
        *,
        number="1",
        sheet="AWAMU_2023_24",
        title="Research title",
        category="Health and Allied Sciences",
        year="2023_2024",
    ):
        return SpecialEventPublication.objects.create(
            participant=participant,
            source_sheet=sheet,
            source_number=number,
            source_row_index=int(number) + 1,
            research_title=title,
            award_category=category,
            award_year=year,
        )

    def test_model_rejects_researcher_for_non_special_event(self):
        participant = SpecialEventParticipant(
            event=self.other_event,
            full_name="Wrong Category",
        )
        with self.assertRaises(ValidationError):
            participant.full_clean()

    def test_excel_import_keeps_same_names_at_different_institutions_separate(self):
        result = import_special_event_participants(
            event=self.event,
            uploaded_file=self.workbook_upload(),
            user=self.user,
        )
        self.assertEqual(result.researchers_created, 3)
        self.assertEqual(result.publications_created, 3)
        self.assertEqual(result.sheets, 2)
        participants = SpecialEventParticipant.objects.filter(event=self.event)
        self.assertEqual(participants.count(), 3)
        duplicated_names = participants.filter(full_name="Repeated Name")
        self.assertEqual(duplicated_names.count(), 2)
        self.assertEqual(
            len(set(duplicated_names.values_list("verification_token", flat=True))),
            2,
        )

    def test_excel_import_groups_publications_under_one_researcher_qr(self):
        result = import_special_event_participants(
            event=self.event,
            uploaded_file=self.workbook_upload({
                "AWAMU_2023_24": [
                    (1, "Dr. Multiple", "University A", "First title", "Health", "2023_2024"),
                    (2, "  Dr. Multiple  ", "University   A", "Second title", "Agriculture", "2023_2024"),
                ],
            }),
            user=self.user,
        )
        self.assertEqual(result.researchers_created, 1)
        self.assertEqual(result.publications_created, 2)
        participant = SpecialEventParticipant.objects.get()
        self.assertEqual(participant.publications.filter(is_active=True).count(), 2)

    def test_reimport_preserves_qr_and_synchronizes_publications(self):
        import_special_event_participants(
            event=self.event,
            uploaded_file=self.workbook_upload({
                "AWAMU_2023_24": [
                    (1, "Researcher", "Institution A", "Original title", "Health", "2023_2024"),
                ],
            }),
            user=self.user,
        )
        original = SpecialEventParticipant.objects.get()
        original_token = original.verification_token
        result = import_special_event_participants(
            event=self.event,
            uploaded_file=self.workbook_upload({
                "AWAMU_2023_24": [
                    (2, "Researcher", "Institution A", "Replacement title", "Engineering", "2023_2024"),
                ],
            }),
            user=self.user,
        )
        original.refresh_from_db()
        self.assertEqual(result.researchers_updated, 1)
        self.assertEqual(SpecialEventParticipant.objects.count(), 1)
        self.assertEqual(original.verification_token, original_token)
        self.assertEqual(original.publications.filter(is_active=True).count(), 1)
        self.assertEqual(
            original.publications.get(is_active=True).research_title,
            "Replacement title",
        )
        self.assertEqual(original.publications.filter(is_active=False).count(), 1)
        self.client.force_login(self.user)
        list_url = reverse("events:special_event_participant_list")
        response = self.client.get(
            f"{list_url}?event={self.event.pk}&q=Original%20title"
        )
        self.assertContains(response, "0</strong><span>researcher QR records")

    def test_public_scan_displays_one_publication_in_english(self):
        participant = self.create_researcher(
            name="Correct Researcher",
            institution="Correct Institution",
        )
        self.create_publication(
            participant,
            number="4",
            title="Correct research title",
            category="Health and Allied Sciences",
            year="2023_2024",
        )
        response = self.client.get(reverse(
            "events:special_event_participant_verify",
            kwargs={"token": participant.verification_token},
        ))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Name:")
        self.assertContains(response, "Correct Researcher")
        self.assertContains(response, "Institution:")
        self.assertContains(response, "Research Title:")
        self.assertContains(response, "Award Category:")
        self.assertContains(response, "Award Year:")
        self.assertContains(response, "© MoEST 2026")
        self.assertNotContains(response, 'class="qr-publication-records"')
        self.assertNotContains(response, "Jina la Mtafiti")
        self.assertNotContains(response, "Taasisi:")
        self.assertNotContains(response, "Nyanja:")
        self.assertNotContains(response, "Registration Status")
        self.assertNotContains(response, "Phase / worksheet")
        self.assertNotContains(response, "Row number")

    def test_public_scan_numbers_all_publications_for_one_researcher(self):
        participant = self.create_researcher(
            name="Dr. Four Awards",
            institution="University A",
        )
        for number in range(1, 5):
            self.create_publication(
                participant,
                number=str(number),
                title=f"Research title {number}",
                category=f"Award category {number}",
                year=f"202{number}_202{number + 1}",
            )
        response = self.client.get(reverse(
            "events:special_event_participant_verify",
            kwargs={"token": participant.verification_token},
        ))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '<ol class="qr-publication-records">')
        self.assertContains(response, "Research Title:", count=4)
        self.assertContains(response, "Award Category:", count=4)
        self.assertContains(response, "Award Year:", count=4)
        self.assertContains(response, "Research title 1")
        self.assertContains(response, "Research title 4")

    def test_qr_endpoint_returns_png_without_numeric_record_id_in_url(self):
        participant = self.create_researcher(name="QR Researcher")
        url = reverse(
            "events:special_event_participant_qr",
            kwargs={"token": participant.verification_token},
        )
        self.assertIn(str(participant.verification_token), url)
        self.assertNotIn(f"/{participant.pk}/", url)
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], "image/png")
        self.assertTrue(response.content.startswith(b"\x89PNG"))

    def test_participant_card_download_returns_paste_ready_png(self):
        participant = self.create_researcher(name="Card Researcher")
        self.create_publication(participant)
        response = self.client.get(reverse(
            "events:special_event_participant_card_download",
            kwargs={"token": participant.verification_token},
        ))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], "image/png")
        self.assertIn("attachment;", response["Content-Disposition"])
        image = Image.open(BytesIO(response.content))
        self.assertEqual(image.format, "PNG")
        self.assertEqual(image.size, (1600, 760))

    def test_qr_and_text_sections_download_as_separate_png_files(self):
        participant = self.create_researcher(name="Separate Image Researcher")
        self.create_publication(participant)
        downloads = (
            ("events:special_event_participant_qr_download", (1000, 1000)),
            ("events:special_event_participant_text_download", (1200, 560)),
        )
        for route_name, expected_size in downloads:
            with self.subTest(route_name=route_name):
                response = self.client.get(reverse(
                    route_name,
                    kwargs={"token": participant.verification_token},
                ))
                self.assertEqual(response.status_code, 200)
                self.assertEqual(response["Content-Type"], "image/png")
                self.assertIn("attachment;", response["Content-Disposition"])
                image = Image.open(BytesIO(response.content))
                self.assertEqual(image.format, "PNG")
                self.assertEqual(image.size, expected_size)

    def test_staff_can_download_all_researcher_cards_as_zip(self):
        for number in ("1", "2"):
            participant = self.create_researcher(
                name=f"Researcher {number}",
                institution="Institution",
            )
            self.create_publication(participant, number=number)
        url = reverse("events:special_event_participant_cards_zip")
        self.assertEqual(self.client.get(f"{url}?event={self.event.pk}").status_code, 302)

        self.client.force_login(self.user)
        response = self.client.get(f"{url}?event={self.event.pk}")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], "application/zip")
        self.assertIn("attachment;", response["Content-Disposition"])
        with ZipFile(BytesIO(response.content)) as archive:
            card_names = [name for name in archive.namelist() if name.endswith(".png")]
            self.assertEqual(len(card_names), 6)
            self.assertEqual(len([name for name in card_names if "/qr-only/" in name]), 2)
            self.assertEqual(len([name for name in card_names if "/text-only/" in name]), 2)
            self.assertEqual(len([name for name in card_names if "/combined-cards/" in name]), 2)
            self.assertTrue(archive.read(card_names[0]).startswith(b"\x89PNG"))

    def test_staff_can_download_one_editable_word_file_with_one_qr_per_researcher(self):
        for number in ("1", "2", "3"):
            participant = self.create_researcher(
                name=f"Word Researcher {number}",
                institution=f"Institution {number}",
            )
            self.create_publication(participant, number=number)
            if number == "1":
                self.create_publication(
                    participant,
                    number="11",
                    title="Second publication for researcher one",
                )
        url = reverse("events:special_event_participant_cards_word")
        self.assertEqual(self.client.get(f"{url}?event={self.event.pk}").status_code, 302)

        self.client.force_login(self.user)
        response = self.client.get(f"{url}?event={self.event.pk}")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("attachment;", response["Content-Disposition"])
        document = Document(BytesIO(response.content))
        self.assertEqual(len(document.inline_shapes), 3)
        self.assertEqual(len(document.tables), 3)
        document_text = "\n".join(
            paragraph.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
        )
        self.assertIn("Word Researcher 1", document_text)
        self.assertIn("2 publications", document_text)
        self.assertIn("Word Researcher 3", document_text)

    def test_staff_list_and_import_require_events_permissions(self):
        list_url = reverse("events:special_event_participant_list")
        import_url = reverse("events:special_event_participant_import")
        self.assertEqual(self.client.get(list_url).status_code, 302)

        participant_user = User.objects.create_user(
            username="ordinary.participant",
            email="ordinary-participant@example.test",
            password=self.password,
            role=User.Role.PARTICIPANT,
        )
        self.client.force_login(participant_user)
        self.assertNotEqual(self.client.get(list_url).status_code, 200)
        self.assertNotEqual(self.client.get(import_url).status_code, 200)

        self.client.force_login(self.user)
        participant = self.create_researcher(name="Download Researcher")
        self.create_publication(participant)
        response = self.client.get(f"{list_url}?event={self.event.pk}")
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Researcher QR records")
        self.assertContains(response, "Download editable Word cards")
        self.assertContains(response, "Research title")
        self.assertNotContains(response, "Jina la Mtafiti")
        self.assertNotContains(response, "Download QR only")
        self.assertNotContains(response, "Download text only")

    def test_import_view_rejects_non_special_event(self):
        self.client.force_login(self.user)
        response = self.client.post(
            reverse("events:special_event_participant_import"),
            {"event": self.other_event.pk, "workbook": self.workbook_upload()},
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Select a valid choice")
        self.assertFalse(SpecialEventParticipant.objects.exists())
